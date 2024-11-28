import os
from PyPDF2 import PdfReader, PdfWriter
from flask import Flask, request, render_template, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
from PIL import Image
import pandas as pd 
from pyPDF2 import pdfReader, pdfWriter
from docx import Document
from reportlab.pdfgen import canvas


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['SECRET_KEY'] = 'your_secret_key'


os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)


ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'jpg', 'png', 'pdf', 'txt', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)

    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        
        target_format = request.form['target_format'].lower()
        base_filename, _ = os.path.splitext(filename)
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{base_filename}.{target_format}")

        try:
           
            if target_format in ['jpg', 'png']:
                img = Image.open(input_path)
                img.convert('RGB').save(output_path, target_format.upper())
            elif target_format in ['csv', 'xlsx']:
                df = pd.read_csv(input_path) if filename.endswith('.csv') else pd.read_excel(input_path)
                if target_format == 'csv':
                    df.to_csv(output_path, index=False)
                else:
                    df.to_excel(output_path, index=False)
            elif target_format == 'txt':
                if filename.endswith('.pdf'):
                    reader = PdfReader(input_path)
                    with open(output_path, 'w') as f:
                        for page in reader.pages:
                            f.write(page.extract_text())
                else:
                    with open(input_path, 'r') as f, open(output_path, 'w') as out:
                        out.write(f.read())
            elif target_format == 'pdf':
                if filename.endswith('.docx'):
                    convert_docx_to_pdf(input_path, output_path)
                else:
                    writer = PdfWriter()
                    writer.add_page(PdfReader(input_path).pages[0]) 
                    with open(output_path, 'wb') as f:
                        writer.write(f)

            
            return send_file(output_path, as_attachment=True)
        except Exception as e:
            flash(f"Error during conversion: {str(e)}")
            return redirect(url_for('index'))

    flash('Invalid file type or conversion format')
    return redirect(url_for('index'))

def convert_docx_to_pdf(input_path, output_path):
   
    document = Document(input_path)
    pdf = canvas.Canvas(output_path)
    flash('Invalid file type or conversion format')
    return redirect(url_for('index'))

def convert_docx_to_pdf(input_path, output_path):
    """
    converts from DOCX to PDF
    """

    document =Document(input_path)
    pdf = canvas.Canvas(output_path)

    for paragraph in document.paragraphs:
        text = paragraph.text
        pdf.drawString(100, pdf._pagesize[1] - 50, text)
        pdf.showPage()  

    pdf.save()

if __name__ == '__main__':
    app.run(debug=True)


